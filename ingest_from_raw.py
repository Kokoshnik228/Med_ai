#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAW (PDF/DOCX/TXT) -> data/*.pages.jsonl + data/manifest.json

Обновления:
- Удаление литературы не «сносит» весь документ, если она встретилась в начале.
- Детектор reference-страниц (по строкам-цитатам) + мягкое удаление «хвоста».
- На смешанных страницах вырезаются только строки-цитаты.
- Устойчивость к OCR-«каше» и частым шаблонам ссылок.

Зависимости (в контейнере app):
  pip install chardet pymupdf pillow python-docx
  # для OCR (опционально):
  apt-get install -y tesseract-ocr tesseract-ocr-eng tesseract-ocr-rus
  pip install easyocr opencv-python-headless
"""

from __future__ import annotations
import argparse
import json
import sys
import os
import re
import hashlib
from pathlib import Path
from typing import Dict, Any, List, Tuple
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing as mp
from time import perf_counter

import chardet
import numpy as np

# ---------- опциональные импорты ----------
try:
    import fitz  # PyMuPDF
except Exception as e:
    print("[ERR] Требуется PyMuPDF: pip install pymupdf", file=sys.stderr, flush=True)
    raise

try:
    from docx import Document  # python-docx==1.1.2
except Exception:
    Document = None

try:
    import pytesseract
    from PIL import Image
    TESS_AVAILABLE = True
except Exception:
    from PIL import Image
    TESS_AVAILABLE = False

try:
    import easyocr
    EASY_AVAILABLE = True
except Exception:
    EASY_AVAILABLE = False

try:
    import cv2
    CV_AVAILABLE = True
except Exception:
    CV_AVAILABLE = False

try:
    import torch
    _TORCH_OK = True
except Exception:
    torch = None
    _TORCH_OK = False


# ================== Параметры и пороги ==================

def _env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except Exception:
        return default

def _env_truthy(name: str, default: bool) -> bool:
    v = os.getenv(name, None)
    if v is None:
        return default
    return v.strip().lower() in ("1","true","yes","on")

# лимиты выводимых фрагментов (для безопасной обрезки больших страниц)
DEFAULT_PAGE_SIZE_CHARS = _env_int("PAGE_SIZE_CHARS", 1800)

# --- детекция литературы ---
REFS_MIN_LINES_PAGE   = _env_int("REFS_MIN_LINES_PAGE", 5)      # мин. строк на странице, чтобы признать её литературной
REFS_RATIO_PAGE       = float(os.getenv("REFS_RATIO_PAGE", "0.55"))  # доля строк-цитат на странице
REFS_CONSEC_MIN       = _env_int("REFS_CONSEC_MIN", 2)          # подряд идущих референс-страниц, чтобы обрезать хвост
REFS_TAIL_FRACTION    = float(os.getenv("REFS_TAIL_FRACTION", "0.6")) # «ближе к концу» — после этой доли страниц
REFS_MIN_PAGES_FROM_START = _env_int("REFS_MIN_PAGES_FROM_START", 3)  # не обрезать всё, если референсы начались слишком рано
REFS_MIN_TOTAL_CHARS  = _env_int("REFS_MIN_TOTAL_CHARS", 5000)  # общее накопленное кол-во символов до среза хвоста

# ================== Утилиты ==================

def file_sha1(p: Path) -> str:
    h = hashlib.sha1()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)

def detect_text_file(path: Path) -> str:
    data = path.read_bytes()
    enc = chardet.detect(data).get("encoding") or "utf-8"
    try:
        return data.decode(enc, errors="ignore")
    except Exception:
        return data.decode("utf-8", errors="ignore")

def pixmap_to_pil(pix: "fitz.Pixmap") -> "Image.Image":
    try:
        needs_colorspace_convert = getattr(pix.colorspace, "n", 3) != 3
    except Exception:
        needs_colorspace_convert = False
    if pix.alpha or needs_colorspace_convert:
        pix = fitz.Pixmap(fitz.csRGB, pix)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img

# --- мягкая нормализация символов (частые OCR-подмены латиница<->кириллица) ---
_LATIN_TO_CYR = str.maketrans({
    "A":"А","a":"а","B":"В","E":"Е","e":"е","K":"К","k":"к","M":"М","H":"Н","O":"О","o":"о",
    "P":"Р","p":"р","C":"С","c":"с","T":"Т","X":"Х","x":"х","Y":"У","y":"у"
})

def clean_text(text: str) -> str:
    t = (text or "").replace("\r", "")
    t = re.sub(r"-\n", "", t)
    t = t.replace("\n\n", "<<<PARA>>>").replace("\n", " ").replace("<<<PARA>>>", "\n\n")
    t = re.sub(r"[ \t]+", " ", t)
    t = t.translate(_LATIN_TO_CYR)
    return t.strip()

# ================== Детекция «литературы» ==================

# заголовки «Список литературы / References» (в т.ч. шумные)
REFS_HDR_RE = re.compile(
    r'^\s*(?:'
    r'спис[оo]к\s+литератур[аы]|литератур[аы]|источники|'
    r'использованн[а-яё]+\s+литератур[аы]|'
    r'references?|bibliograph\w*'
    r')\s*[:\-–—]?\s*$',
    re.IGNORECASE | re.MULTILINE
)

# строка похожа на библиографическую ссылку?
CITATION_LINE_RE = re.compile(
    r'''(?xi)
    (?:^\s*\d{1,3}[\).\]]\s+)                       # нумерация пункта
    | (?:\b(et\s*al\.?|еt\s*аl\.?)\b)               # et al.
    | (?:\bdoi[:\s/]|10\.\d{3,9}/\S+)               # doi
    | (?:\bEpub\b|\bPublished\b|\bRetrieved\b)      # пометки публикации
    | (?:\b\d{4}\b\s*;\s*\d{1,4}\s*(?:\(\d{1,4}\))?\s*:\s*\d{1,5}(?:[-–]\d{1,5})?) # 2019;54(6):1157-1170
    ''',
    re.UNICODE
)

def is_citation_line(line: str) -> bool:
    s = (line or "").strip()
    if not s:
        return False
    # короткие огрызки — не считаем
    if len(s) < 20:
        return False
    return bool(CITATION_LINE_RE.search(s))

def split_lines_keep(text: str) -> List[str]:
    # мягко разбиваем на «строки» для анализа ссылок
    return [x.strip() for x in re.split(r'(?:\n+|(?<=\.)\s+)', text or "") if x and x.strip()]

def classify_references_page(text: str) -> Tuple[bool, float, int]:
    """Возвращает (is_refs_page, ratio, n_lines)"""
    lines = split_lines_keep(text)
    if len(lines) < REFS_MIN_LINES_PAGE:
        return (False, 0.0, len(lines))
    hits = sum(1 for ln in lines if is_citation_line(ln))
    ratio = hits / max(len(lines), 1)
    return (ratio >= REFS_RATIO_PAGE, ratio, len(lines))

def drop_citation_lines(text: str) -> str:
    """Удаляет строки, похожие на библиографические записи, оставляет остальной текст."""
    lines = split_lines_keep(text)
    kept = [ln for ln in lines if not is_citation_line(ln)]
    out = "\n".join(kept).strip()
    return out

def decide_tail_cut(page_flags: List[bool], total_pages: int, total_chars_before: List[int]) -> int:
    """
    Если ближе к концу идёт блок из ≥REFS_CONSEC_MIN референс-страниц, вернём
    индекс первой такой страницы для «хвостового» обрезания. Иначе -1.
    total_chars_before[i] — накопленные символы ДО страницы i (0-based).
    """
    if total_pages == 0:
        return -1
    start_tail = int(total_pages * REFS_TAIL_FRACTION)
    consec = 0
    first_idx = -1
    for i in range(total_pages):
        if not page_flags[i]:
            consec = 0
            first_idx = -1
            continue
        # считаем только «ближе к концу»
        if i < start_tail:
            continue
        consec = consec + 1 if consec > 0 else 1
        if consec == 1:
            first_idx = i
        if consec >= REFS_CONSEC_MIN and total_chars_before[i] >= REFS_MIN_TOTAL_CHARS:
            return max(first_idx, 0)
    return -1

# ================== DOCX/TXT ==================

def extract_docx_text(path: Path) -> str:
    if Document is None:
        print("[ERR] python-docx не установлен. Добавь `python-docx` в зависимости.", file=sys.stderr, flush=True)
        return ""
    try:
        doc = Document(str(path))
    except Exception as e:
        print(f"[ERR] Не удалось открыть DOCX {path.name}: {e}", file=sys.stderr, flush=True)
        return ""
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [ (c.text or "").strip() for c in row.cells ]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)
    return "\n".join(parts).strip()

def split_text_to_pages(full_text: str, page_size_chars: int = DEFAULT_PAGE_SIZE_CHARS) -> List[Dict[str, Any]]:
    text = clean_text(full_text)
    if not text:
        return [{"page": 1, "text": ""}]

    # разрежем примерно по абзацам/предложениям с ограничением размера
    parts: List[str] = []
    buf = []
    cur_len = 0
    tokens = re.split(r"(\n\n|[.!?]\s+)", text)
    for t in tokens:
        if t is None:
            continue
        if cur_len + len(t) > page_size_chars and buf:
            parts.append("".join(buf).strip())
            buf, cur_len = [t], len(t)
        else:
            buf.append(t)
            cur_len += len(t)
    if buf:
        parts.append("".join(buf).strip())

    pages = []
    for i, chunk in enumerate(parts, start=1):
        pages.append({"page": i, "text": chunk})
    return pages or [{"page": 1, "text": text[:page_size_chars]}]

# ================== OCR Backends ==================

def ocr_page_tesseract(img_pil: "Image.Image", lang: str) -> str:
    if not TESS_AVAILABLE:
        return ""
    os.environ.setdefault("OMP_NUM_THREADS", "1")
    os.environ.setdefault("OPENBLAS_NUM_THREADS", "1")
    cfg = "--oem 1 --psm 6"
    return (pytesseract.image_to_string(img_pil, lang=lang, config=cfg) or "").strip()

def ocr_page_easyocr(img_pil: "Image.Image", reader) -> str:
    if reader is None:
        return ""
    arr = np.array(img_pil.convert("RGB"))
    res = reader.readtext(arr, detail=0, paragraph=True)
    return "\n".join([x.strip() for x in res if x]).strip()

def preprocess_pil(img_pil: "Image.Image") -> "Image.Image":
    if not CV_AVAILABLE:
        return img_pil
    img = np.array(img_pil.convert("L"))
    try:
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        img = clahe.apply(img)
    except Exception:
        pass
    img = cv2.fastNlMeansDenoising(img, h=10)
    img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                cv2.THRESH_BINARY, 35, 15)
    return Image.fromarray(img)

# ================== Извлечение PDF с анти-литературой ==================

def ingest_pdf(
    pdf_path: Path,
    *,
    ocr_mode: str,            # "auto"|"always"|"never"
    ocr_backend: str,         # "tesseract"|"easyocr"
    ocr_lang: str,
    dpi: int,
    min_chars: int,
    verbose: bool,
    easy_reader=None
) -> List[Dict[str, Any]]:

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"[WARN] Не удалось открыть PDF {pdf_path.name}: {e}", file=sys.stderr, flush=True)
        return []

    raw_pages: List[str] = []
    for i, page in enumerate(doc, start=1):
        txt = (page.get_text("text") or "").strip()

        # решаем, делать ли OCR
        if ocr_mode == "always":
            do_ocr = True
        elif ocr_mode == "auto":
            do_ocr = (len(txt) < min_chars)
        else:
            do_ocr = False

        if do_ocr:
            if ocr_backend == "tesseract":
                zoom = dpi / 72.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_pil = pixmap_to_pil(pix)
                img_pil = preprocess_pil(img_pil)
                txt_ocr = ocr_page_tesseract(img_pil, ocr_lang)
            elif ocr_backend == "easyocr":
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                img_pil = pixmap_to_pil(pix)
                img_pil = preprocess_pil(img_pil)
                txt_ocr = ocr_page_easyocr(img_pil, easy_reader)
            else:
                txt_ocr = ""
            if len(txt_ocr) > len(txt):
                txt = txt_ocr
                if verbose:
                    print(f"[OCR-{ocr_backend}] {pdf_path.name} p.{i}: len={len(txt)}", flush=True)

        raw_pages.append(txt)

    # --- глобальная обработка литературы ---
    total_pages = len(raw_pages)
    cleaned_pages: List[str] = []
    refs_flags: List[bool] = []
    total_chars_before: List[int] = []
    acc = 0

    # предварительно чистим, смотрим заголовки и долю ссылок
    header_hits_idx: List[int] = []
    for idx, t in enumerate(raw_pages):
        t0 = t or ""
        # сначала легкая чистка
        t1 = clean_text(t0)
        # заголовок «References»/«Список литературы»?
        if REFS_HDR_RE.search(t1):
            header_hits_idx.append(idx)
        # страница — почти полностью из ссылок?
        is_ref, ratio, n_lines = classify_references_page(t1)
        refs_flags.append(is_ref)
        cleaned_pages.append(t1)
        total_chars_before.append(acc)
        acc += len(t1)

    # решаем, нужно ли отрезать хвост
    cut_from = decide_tail_cut(refs_flags, total_pages, total_chars_before)

    pages_out: List[Dict[str, Any]] = []
    for i, text in enumerate(cleaned_pages, start=1):
        if cut_from >= 0 and (i-1) >= cut_from:
            # «хвост» после блока литературы ближе к концу — полностью пропускаем
            continue

        # если страница выглядит как литература и она слишком рано (в начале) — просто пропустим эту страницу,
        # но документ не трогаем целиком
        early_refs = (i <= REFS_MIN_PAGES_FROM_START) and refs_flags[i-1]
        if early_refs:
            continue

        # если страница «в основном литература», но не попала под ранние или хвостовые правила — удалим строки-ссылки
        if refs_flags[i-1]:
            text = drop_citation_lines(text)

        # и ещё раз локально: если в тексте случайные ссылки — уберём их
        text = drop_citation_lines(text)

        pages_out.append({"page": i, "text": text})

    return pages_out

# ================== TXT/DOCX с анти-литературой ==================

def ingest_txt(txt_path: Path, page_size_chars: int = DEFAULT_PAGE_SIZE_CHARS) -> List[Dict[str, Any]]:
    text = detect_text_file(txt_path).strip()
    text = clean_text(text)
    # порежем на псевдостраницы
    pages = split_text_to_pages(text, page_size_chars=page_size_chars)

    # применим такую же логику, как для PDF
    flags = []
    total_chars_before = []
    acc = 0
    for p in pages:
        t = p.get("text","")
        is_ref, _, _ = classify_references_page(t)
        flags.append(is_ref)
        total_chars_before.append(acc)
        acc += len(t)

    cut_from = decide_tail_cut(flags, len(pages), total_chars_before)

    out: List[Dict[str, Any]] = []
    for i, p in enumerate(pages, start=1):
        if cut_from >= 0 and (i-1) >= cut_from:
            continue
        t = p.get("text","")
        early_refs = (i <= REFS_MIN_PAGES_FROM_START) and flags[i-1]
        if early_refs:
            continue
        if flags[i-1]:
            t = drop_citation_lines(t)
        t = drop_citation_lines(t)
        out.append({"page": i, "text": t})
    return out

def ingest_docx(path: Path, page_size_chars: int = DEFAULT_PAGE_SIZE_CHARS) -> List[Dict[str, Any]]:
    raw = extract_docx_text(path)
    if not raw:
        return [{"page": 1, "text": ""}]
    raw = clean_text(raw)
    return ingest_txt_like(raw, page_size_chars=page_size_chars)

def ingest_txt_like(full_text: str, page_size_chars: int = DEFAULT_PAGE_SIZE_CHARS) -> List[Dict[str, Any]]:
    pages = split_text_to_pages(full_text, page_size_chars=page_size_chars)
    flags = []
    total_chars_before = []
    acc = 0
    for p in pages:
        t = p.get("text","")
        is_ref, _, _ = classify_references_page(t)
        flags.append(is_ref)
        total_chars_before.append(acc)
        acc += len(t)
    cut_from = decide_tail_cut(flags, len(pages), total_chars_before)

    out: List[Dict[str, Any]] = []
    for i, p in enumerate(pages, start=1):
        if cut_from >= 0 and (i-1) >= cut_from:
            continue
        t = p.get("text","")
        early_refs = (i <= REFS_MIN_PAGES_FROM_START) and flags[i-1]
        if early_refs:
            continue
        if flags[i-1]:
            t = drop_citation_lines(t)
        t = drop_citation_lines(t)
        out.append({"page": i, "text": t})
    return out or [{"page": 1, "text": ""}]

# ================== Основной процесс ==================

def choose_ocr_backend(requested: str) -> str:
    req = (requested or "").lower()
    if req == "easyocr" and EASY_AVAILABLE:
        return "easyocr"
    if req == "tesseract" and TESS_AVAILABLE:
        return "tesseract"
    if EASY_AVAILABLE:
        return "easyocr"
    if TESS_AVAILABLE:
        return "tesseract"
    return "none"

def _easyocr_models_ready(model_dir: Path) -> bool:
    try:
        mdir = model_dir / "model"
        return mdir.exists() and any(mdir.iterdir())
    except Exception:
        return False

def process_one_file(
    f: Path,
    out_dir: Path,
    *,
    min_chars: int,
    ocr_mode: str,
    ocr_backend_eff: str,
    ocr_lang: str,
    dpi: int,
    verbose: bool,
    page_size_chars: int,
    easyocr_dir: Path,
    easyocr_use_gpu: bool,
    easyocr_allow_downloads: bool,
) -> Dict[str, Any]:
    """Процессинг одного файла (без доступа к manifest). Возвращает entry + путь pages."""
    sha = file_sha1(f)
    stem = f.stem
    doc_id = stem
    out_pages = out_dir / f"{doc_id}.pages.jsonl"

    easy_reader = None
    if ocr_mode != "never" and ocr_backend_eff == "easyocr":
        easy_reader = easyocr.Reader(
            ['ru', 'en'],
            gpu=easyocr_use_gpu,
            model_storage_directory=str(easyocr_dir),
            download_enabled=bool(easyocr_allow_downloads),
            verbose=False,
        )

    ext = f.suffix.lower()
    if ext == ".pdf":
        pages = ingest_pdf(
            f, ocr_mode=ocr_mode, ocr_backend=ocr_backend_eff, ocr_lang=ocr_lang,
            dpi=dpi, min_chars=min_chars, verbose=verbose, easy_reader=easy_reader
        )
    elif ext == ".docx":
        pages = ingest_docx(f, page_size_chars=page_size_chars)
    else:
        pages = ingest_txt(f, page_size_chars=page_size_chars)

    with out_pages.open("w", encoding="utf-8") as w:
        for p in pages:
            rec = {"doc_id": doc_id, "page": p.get("page", 1), "text": p.get("text", "")}
            w.write(json.dumps(rec, ensure_ascii=False) + "\n")

    empty_pages = sum(1 for p in pages if not (p.get("text") or "").strip())

    return {
        "doc_id": doc_id,
        "source_path": str(f),
        "pages": len(pages),
        "lang": "ru",
        "sha1": sha,
        "ocr_backend": ocr_backend_eff,
        "ocr_mode": ocr_mode,
        "ocr_lang": ocr_lang,
        "dpi": dpi,
        "min_chars": min_chars,
        "empty_pages": empty_pages,
        "out_pages": str(out_pages)
    }

def _decide_easyocr_gpu(ocr_gpu_arg: str) -> bool:
    ocr_gpu_arg = (ocr_gpu_arg or "auto").lower()
    if ocr_gpu_arg == "cuda":
        return _TORCH_OK and torch.cuda.is_available()
    if ocr_gpu_arg == "cpu":
        return False
    return _TORCH_OK and torch.cuda.is_available()

def main():
    ap = argparse.ArgumentParser("RAW -> data/*.pages.jsonl (+manifest) с OCR и анти-литературой")
    ap.add_argument("--input-dir", default="raw_docs", help="Папка с PDF/DOCX/TXT (рекурсивно)")
    ap.add_argument("--out-dir", default="data", help="Куда сохранять JSONL и manifest.json")
    ap.add_argument("--force", action="store_true", help="Перепарсить даже без изменений")

    # OCR
    ap.add_argument("--ocr-mode", choices=["auto","always","never"], default=os.getenv("OCR_MODE", "auto"))
    ap.add_argument("--ocr-backend", choices=["tesseract","easyocr"], default=os.getenv("OCR_BACKEND", "easyocr"))
    ap.add_argument("--ocr-lang", default=os.getenv("TESS_LANG", "rus+eng"))
    ap.add_argument("--min-chars", type=int, default=int(os.getenv("MIN_CHARS", "60")))
    ap.add_argument("--dpi", type=int, default=int(os.getenv("OCR_DPI", "300")))

    # EasyOCR/GPU
    ap.add_argument("--ocr-gpu", choices=["auto","cpu","cuda"], default=os.getenv("OCR_GPU", "auto"))
    ap.add_argument("--easyocr-dir", default=os.getenv("EASYOCR_DIR", str(Path.home() / ".EasyOCR")))
    ap.add_argument("--easyocr-allow-downloads", action="store_true",
                    default=os.getenv("EASYOCR_ALLOW_DOWNLOADS", "0").lower() in ("1","true","yes"))

    # прочее
    ap.add_argument("--workers", type=int, default=0)
    ap.add_argument("--page-size-chars", type=int, default=DEFAULT_PAGE_SIZE_CHARS)
    ap.add_argument("--verbose", action="store_true")

    args = ap.parse_args()

    in_dir = Path(args.input_dir)
    out_dir = Path(args.out_dir)
    ensure_dir(out_dir)

    allowed = {".pdf", ".docx", ".txt"}
    files: List[Path] = []
    for p in in_dir.rglob("*"):
        if not p.is_file():
            continue
        ext = p.suffix.lower()
        if ext not in allowed:
            continue
        name = p.name
        if name.startswith("~$"):
            continue
        files.append(p)
    files = sorted(files)

    if not files:
        print(f"В {in_dir} нет pdf/docx/txt", file=sys.stderr, flush=True)
        return 1

    manifest_path = out_dir / "manifest.json"
    manifest: Dict[str, Any] = {"docs": []}
    if manifest_path.exists():
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            if not isinstance(manifest, dict) or "docs" not in manifest:
                manifest = {"docs": []}
        except Exception:
            manifest = {"docs": []}

    docs = manifest.get("docs", [])
    by_source: Dict[str, Dict[str, Any]] = {d.get("source_path"): d for d in docs if isinstance(d, dict)}
    existing_ids = {d.get("doc_id") for d in docs if isinstance(d, dict)}

    plan: List[Path] = []
    for f in files:
        sha = file_sha1(f)
        prev = by_source.get(str(f))
        if args.force or not prev or prev.get("sha1") != sha:
            plan.append(f)
        elif args.verbose:
            try:
                rel = f.relative_to(in_dir)
            except Exception:
                rel = f
            print(f"→ Без изменений: {rel}", flush=True)

    if not plan:
        print("Нет изменений — ничего делать не нужно.", flush=True)
        return 0

    ocr_backend_eff = choose_ocr_backend(args.ocr_backend)
    if ocr_backend_eff == "none":
        args.ocr_mode = "never"

    easyocr_dir = Path(args.easyocr_dir).expanduser()
    ensure_dir(easyocr_dir / "model")

    need_easy_warmup = (
        args.ocr_mode != "never"
        and ocr_backend_eff == "easyocr"
        and not _easyocr_models_ready(easyocr_dir)
    )
    if need_easy_warmup:
        print("⏳ EasyOCR warmup: загрузка моделей (один раз)...", flush=True)
        use_gpu = _decide_easyocr_gpu(args.ocr_gpu)
        easyocr.Reader(['ru','en'], gpu=use_gpu,
                       model_storage_directory=str(easyocr_dir),
                       download_enabled=bool(args.easyocr_allow_downloads),
                       verbose=False)
        args.workers = 1

    workers = args.workers or max(1, os.cpu_count() or 1)
    easyocr_use_gpu = False
    if args.ocr_mode != "never" and ocr_backend_eff == "easyocr":
        if workers > 1:
            print("⚠️ EasyOCR: переключаюсь на workers=1 для стабильности и GPU.", flush=True)
            workers = 1
        easyocr_use_gpu = _decide_easyocr_gpu(args.ocr_gpu)
        print(f"EasyOCR init планируется с GPU={easyocr_use_gpu}", flush=True)

    total = len(plan)
    print(f"📦 Ingest started: {total} files (workers={workers}, ocr={args.ocr_mode}/{ocr_backend_eff})", flush=True)

    results: List[Dict[str, Any]] = []
    t_start = perf_counter()

    if workers <= 1 or len(plan) == 1:
        for i, f in enumerate(plan, 1):
            t0 = perf_counter()
            try:
                r = process_one_file(
                    f, out_dir,
                    min_chars=args.min_chars,
                    ocr_mode=args.ocr_mode,
                    ocr_backend_eff=ocr_backend_eff,
                    ocr_lang=args.ocr_lang,
                    dpi=args.dpi,
                    verbose=args.verbose,
                    page_size_chars=args.page_size_chars,
                    easyocr_dir=easyocr_dir,
                    easyocr_use_gpu=easyocr_use_gpu,
                    easyocr_allow_downloads=False,
                )
                results.append(r)
                dt = perf_counter() - t0
                print(f"[{i}/{total}] {f.name}: {r['pages']} pages, empty={r['empty_pages']}, ocr={r['ocr_backend']}/{r['ocr_mode']} ({dt:.2f}s)", flush=True)
            except Exception as e:
                print(f"[ERR] {f.name}: {e}", file=sys.stderr, flush=True)
    else:
        mp_ctx = mp.get_context("spawn")
        with ProcessPoolExecutor(max_workers=workers, mp_context=mp_ctx) as ex:
            futs = {
                ex.submit(
                    process_one_file, f, out_dir,
                    min_chars=args.min_chars,
                    ocr_mode=args.ocr_mode,
                    ocr_backend_eff=ocr_backend_eff,
                    ocr_lang=args.ocr_lang,
                    dpi=args.dpi,
                    verbose=args.verbose,
                    page_size_chars=args.page_size_chars,
                    easyocr_dir=easyocr_dir,
                    easyocr_use_gpu=False,
                    easyocr_allow_downloads=False,
                ): f for f in plan
            }
            done = 0
            for fut in as_completed(futs):
                f = futs[fut]
                try:
                    r = fut.result()
                    results.append(r)
                    done += 1
                    print(f"[{done}/{total}] {f.name}: {r['pages']} pages, empty={r['empty_pages']}, ocr={r['ocr_backend']}/{r['ocr_mode']}", flush=True)
                except Exception as e:
                    print(f"[ERR] {f.name}: {e}", file=sys.stderr, flush=True)

    # обновляем manifest и уникализируем doc_id при коллизии
    for r in results:
        src = r["source_path"]
        doc_id = r["doc_id"]

        if doc_id in existing_ids:
            base = doc_id
            suf = 2
            while f"{base}_{suf}" in existing_ids:
                suf += 1
            new_id = f"{base}_{suf}"
            op = Path(r["out_pages"])
            op_renamed = op.with_name(f"{new_id}.pages.jsonl")
            try:
                op.rename(op_renamed)
            except FileNotFoundError:
                pass
            r["doc_id"] = new_id
            r["out_pages"] = str(op_renamed)
            doc_id = new_id
        existing_ids.add(doc_id)

        prev = by_source.get(src)
        entry = {
            "doc_id": doc_id,
            "source_path": src,
            "pages": r["pages"],
            "lang": "ru",
            "sha1": r["sha1"],
            "ocr": (r["ocr_mode"] != "never" and r["ocr_backend"] != "none"),
            "ocr_mode": r["ocr_mode"],
            "ocr_backend": r["ocr_backend"],
            "ocr_lang": r["ocr_lang"],
            "dpi": r["dpi"],
            "min_chars": r["min_chars"],
            "empty_pages": r["empty_pages"],
        }
        if prev:
            prev.update(entry)
        else:
            manifest["docs"].append(entry)

    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    total_dt = perf_counter() - t_start
    print(f"\nГотово ✅: обработано файлов = {len(results)} за {total_dt:.2f}s. Обновлён {manifest_path}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
